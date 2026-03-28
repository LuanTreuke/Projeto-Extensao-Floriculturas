from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Caminho do arquivo de saída
output_path = '../TCC_Interface_Report.docx'

doc = Document()

# Helper para parágrafos de corpo
def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    return p

# Helper para blocos de código (monoespaçado)
def add_code_block(code_text):
    p = doc.add_paragraph()
    r = p.add_run(code_text)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    return p

# Título
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = title.add_run('Desenvolvimento da Interface Inicial — Relatório')
run.bold = True
run.font.size = Pt(14)

add_paragraph('\n1. Desenvolvimento da Interface Inicial', bold=True)
add_paragraph(
    'A interface inicial foi desenvolvida como uma aplicação web utilizando Next.js e React com TypeScript. ' 
    'O projeto organiza rotas e componentes reutilizáveis, incluindo um painel administrativo para gerenciar o catálogo de produtos. ' 
    'A comunicação com o backend é realizada via HTTP usando a biblioteca axios, configurada em um cliente central (`services/api.ts`). ' 
    'A navegação utiliza rotas do Next.js: a listagem de produtos encontra-se em /admin/catalogo, o formulário de adição em /admin/catalogo/adicionar, e a edição em /admin/catalogo/[id]/editar. '
)

add_paragraph('\nTrecho representativo (configuração do cliente HTTP):', bold=True)
add_code_block(
    """// services/api.ts (excerto)
const api = axios.create({ baseURL: API_URL, timeout: 10000 });
api.interceptors.request.use((config) => {
  const raw = localStorage.getItem('usuario');
  if (raw) {
    const u = JSON.parse(raw);
    if (u && (u.id || u.id === 0)) config.headers = { ...(config.headers || {}), ['x-user-id']: String(u.id) };
  }
  return config;
});
"""
)

add_paragraph('\n2. Estrutura da Camada de Dados', bold=True)
add_paragraph(
    'A camada de dados do sistema foi implementada no backend com NestJS e TypeORM. As entidades representam tabelas relacionais e as operações de persistência são realizadas por repositórios TypeORM. ' 
    'A entidade Produto inclui campos para identificação, nome, descrição, preço, URLs de imagem e um campo booleano `enabled`. Existe uma relação ManyToOne com a entidade Categoria.'
)

add_paragraph('\nTrecho representativo (entidade Produto):', bold=True)
add_code_block(
    """// back-tcc-floricultura/src/produto/produto.entity.ts
@Entity('produto')
export class Produto {
  @PrimaryGeneratedColumn() id: number;
  @Column({ length: 100 }) nome: string;
  @Column('text', { nullable: true }) descricao: string;
  @Column('decimal', { precision: 10, scale: 2 }) preco: number;
  @Column('text', { nullable: true }) imagem_url: string;
  @Column({ type: 'boolean', default: true }) enabled: boolean;
  @ManyToOne(() => Categoria, { nullable: true }) @JoinColumn({ name: 'Categoria_id' }) categoria: Categoria;
}
"""
)

add_paragraph('\nTrecho representativo (controle de acesso no controller):', bold=True)
add_code_block(
    """// back-tcc-floricultura/src/produto/produto.controller.ts (excerto)
@Post()
async create(@Headers('x-user-id') userId: string, @Body() dto: CreateProdutoDto) {
  const u = userId ? await this.usuarioService.findById(Number(userId)) : null;
  if (!u || u.role !== 'Admin') throw new ForbiddenException('Acesso negado');
  return this.produtoService.create(dto);
}
"""
)

add_paragraph('\n3. Implementação do Cadastro de Produtos', bold=True)
add_paragraph(
    'O cadastro de produtos no front-end é realizado por meio de um formulário controlado em `/app/admin/catalogo/adicionar/page.tsx`. O fluxo inclui seleção múltipla de imagens por um componente dedicado, upload das imagens via `uploadService`, montagem de um DTO e chamada ao serviço `addProduct` definido em `services/productService.ts`. ' 
    'As URLs de imagem são concatenadas em uma única string separada por vírgulas e armazenadas no campo `imagem_url`.'
)

add_paragraph('\nTrecho representativo (envio do formulário):', bold=True)
add_code_block(
    """// app/admin/catalogo/adicionar/page.tsx (excerto)
const imagemUrls: string[] = [];
for (let i = 0; i < selectedFiles.length; i++) {
  const uploadResult = await uploadImage(selectedFiles[i]);
  imagemUrls.push(uploadResult.url);
}
const dto = { nome, descricao, preco: Number(preco), imagem_url: imagemUrls.join(','), Categoria_id: Number(categoria || 0), enabled };
await addProduct(dto as Omit<Product, 'id'>);
"""
)

add_paragraph('\n4. Implementação da Listagem de Produtos', bold=True)
add_paragraph(
    'A listagem de produtos é implementada na página administrativa `/app/admin/catalogo/page.tsx`. A página carrega os itens via `fetchProducts()` e exibe-os em um grid usando o componente `ProductCard`. ' 
    'Filtros (busca por nome com normalização de acentos, categoria, faixa de preço) e ordenação são aplicados no cliente para melhor responsividade. A interface também oferece ações administrativas: editar, excluir e alternar o estado `enabled` dos produtos, com atualizações otimistas e rollback em caso de falha.'
)

add_paragraph('\nTrecho representativo (busca de produtos):', bold=True)
add_code_block(
    """// services/productService.ts (excerto)
export async function fetchProducts(): Promise<Product[]> {
  const res = await api.get(`/produtos`);
  const data = res.data as unknown;
  if (Array.isArray(data)) return data as Product[];
  if (data && typeof data === 'object' && Array.isArray((data as any).items)) return (data as any).items as Product[];
  return [];
}
"""
)

add_paragraph('\nTrecho representativo (filtragem e renderização):', bold=True)
add_code_block(
    """// app/admin/catalogo/page.tsx (excerto)
{products
  .filter(p => nomeNormalized.includes(searchNormalized) && /* categoria e preço */)
  .sort((a,b) => sort === 'new' ? b.id - a.id : sort === 'asc' ? a.preco - b.preco : b.preco - a.preco)
  .map(p => <ProductCard key={p.id} id={p.id} name={p.nome} price={`R$${Number(p.preco).toFixed(2)}`} image={p.imagem_url?.split(',')[0].trim() || ''} />)}
"""
)

add_paragraph('\n5. Testes Iniciais do Sistema', bold=True)
add_paragraph(
    'Foram realizados testes funcionais iniciais para validar a integração entre front-end e back-end e o fluxo principal do sistema. Entre as verificações: teste de conexão com a API por meio do componente `ApiTest`, testes manuais de inclusão/edição/exclusão de produtos, e verificação de permissões (uso do header `x-user-id` para ações administrativas). O repositório contém infraestrutura de testes que pode ser aproveitada posteriormente.'
)

add_paragraph('\nTrecho representativo (componente de teste API):', bold=True)
add_code_block(
    """// components/ApiTest.tsx (excerto)
const response = await api.get('/produtos');
setTestResult(`✅ API OK (${response.data.length} produtos)`);
"""
)

add_paragraph('\n6. Prints das Telas', bold=True)
add_paragraph('\nFigura 1 – Tela de cadastro de produtos')
doc.add_paragraph('[INSERIR PRINT DA TELA]')
doc.add_page_break()
add_paragraph('\nFigura 2 – Tela de listagem de produtos')
doc.add_paragraph('[INSERIR PRINT DA TELA]')
doc.add_page_break()
add_paragraph('\nFigura 3 – Interface inicial do sistema')
doc.add_paragraph('[INSERIR PRINT DA TELA]')

add_paragraph('\nConsiderações finais', bold=True)
add_paragraph(
    'O sistema é descrito como uma solução digital aplicável a pequenos empreendimentos; a implementação utiliza Next.js + React no front-end, `axios` para comunicação HTTP e NestJS + TypeORM no backend. Os trechos incluídos são reais e extraídos do código-fonte do projeto.'
)

# Salvar documento
with open(output_path, 'wb') as f:
    doc.save(f)

print('DOCX gerado em:', output_path)
